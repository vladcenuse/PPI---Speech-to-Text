from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
import json
import os
from docx import Document
from docx.shared import Inches

app = FastAPI(title="Medical Voice-to-Text Assistant", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Romanian Medical Patient File Model
class PatientFile(BaseModel):
    id: Optional[int] = None
    patient_name: str
    patient_id: str
    date: datetime
    doctor_name: str = "Dr. Alina Baciu"
    
    # Medical sections (highlighted fields to be auto-filled)
    anamneza: Optional[str] = None  # Medical history
    examen_clinic: Optional[str] = None  # Clinical examination
    diagnostic: Optional[str] = None  # Diagnosis
    tratament: Optional[str] = None  # Treatment
    recomandari: Optional[str] = None  # Recommendations
    observatii: Optional[str] = None  # Observations
    
    # Audio and transcription info
    audio_file_path: Optional[str] = None
    transcription: Optional[str] = None
    transcription_status: str = "pending"  # pending, completed, error
    
    created_at: datetime = datetime.now()
    updated_at: datetime = datetime.now()

# In-memory storage for MVP
patient_files_db: List[PatientFile] = []
next_id = 1

# Mock Romanian speech-to-text responses
MOCK_ROMANIAN_TRANSCRIPTIONS = [
    {
        "anamneza": "Pacientul se prezintă cu dureri abdominale persistente de 3 zile, localizate în regiunea epigastrică.",
        "examen_clinic": "Abdomenul este sensibil la palpare, fără defense musculară. Ficatul și splina nu sunt palpabile.",
        "diagnostic": "Gastrită acută, suspiciune de ulcer gastric.",
        "tratament": "Omeprazol 20mg dimineața, ranitidină 150mg seara, dietă blandă.",
        "recomandari": "Evitarea alimentelor picante și acide, control în 2 săptămâni.",
        "observatii": "Pacientul este anxios, necesită explicații suplimentare despre tratament."
    },
    {
        "anamneza": "Femeie de 45 ani, se plânge de cefalee intensă și greață de 2 zile.",
        "examen_clinic": "Tensiune arterială 160/95 mmHg, puls regulat. Fundul de ochi normal.",
        "diagnostic": "Hipertensiune arterială moderată.",
        "tratament": "Amlodipină 5mg zilnic, modificări în stilul de viață.",
        "recomandari": "Măsurarea zilnică a tensiunii, reducerea consumului de sare.",
        "observatii": "Pacientul fumează 10 țigări/zi, necesită consiliere pentru renunțarea la fumat."
    }
]

@app.get("/")
def read_root():
    return {"message": "Medical Voice-to-Text Assistant - Dr. Alina Baciu", "version": "1.0.0"}

@app.get("/api/patient-files")
def get_patient_files():
    """Get all patient files"""
    return patient_files_db

@app.get("/api/patient-files/{file_id}")
def get_patient_file(file_id: int):
    """Get specific patient file"""
    for file in patient_files_db:
        if file.id == file_id:
            return file
    raise HTTPException(status_code=404, detail="Patient file not found")

@app.post("/api/patient-files")
def create_patient_file(patient_file: PatientFile):
    """Create new patient file"""
    global next_id
    patient_file.id = next_id
    next_id += 1
    patient_file.created_at = datetime.now()
    patient_file.updated_at = datetime.now()
    patient_files_db.append(patient_file)
    return patient_file

@app.put("/api/patient-files/{file_id}")
def update_patient_file(file_id: int, patient_file: PatientFile):
    """Update patient file"""
    for i, file in enumerate(patient_files_db):
        if file.id == file_id:
            patient_file.id = file_id
            patient_file.created_at = file.created_at
            patient_file.updated_at = datetime.now()
            patient_files_db[i] = patient_file
            return patient_file
    raise HTTPException(status_code=404, detail="Patient file not found")

@app.delete("/api/patient-files/{file_id}")
def delete_patient_file(file_id: int):
    """Delete patient file"""
    for i, file in enumerate(patient_files_db):
        if file.id == file_id:
            deleted_file = patient_files_db.pop(i)
            return {"message": "Patient file deleted", "file": deleted_file}
    raise HTTPException(status_code=404, detail="Patient file not found")

@app.post("/api/transcribe-audio")
async def transcribe_audio(file: UploadFile = File(...)):
    """Mock Romanian speech-to-text transcription"""
    # Simulate processing time
    import asyncio
    await asyncio.sleep(1)
    
    # Return mock Romanian medical data
    import random
    mock_data = random.choice(MOCK_ROMANIAN_TRANSCRIPTIONS)
    
    return {
        "transcription": "Transcrierea audio a fost procesată cu succes.",
        "extracted_data": mock_data,
        "status": "completed"
    }

@app.post("/api/export-to-word/{file_id}")
def export_to_word(file_id: int):
    """Export patient file to Word document"""
    # Find the patient file
    patient_file = None
    for file in patient_files_db:
        if file.id == file_id:
            patient_file = file
            break
    
    if not patient_file:
        raise HTTPException(status_code=404, detail="Patient file not found")
    
    # Create Word document
    doc = Document()
    doc.add_heading('Fișa Pacientului', 0)
    
    # Add patient info
    doc.add_heading('Informații Pacient', level=1)
    doc.add_paragraph(f'Nume: {patient_file.patient_name}')
    doc.add_paragraph(f'ID Pacient: {patient_file.patient_id}')
    doc.add_paragraph(f'Data: {patient_file.date.strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Medic: {patient_file.doctor_name}')
    
    # Add medical sections
    doc.add_heading('Anamneză', level=1)
    doc.add_paragraph(patient_file.anamneza or 'Nu este completat')
    
    doc.add_heading('Examen Clinic', level=1)
    doc.add_paragraph(patient_file.examen_clinic or 'Nu este completat')
    
    doc.add_heading('Diagnostic', level=1)
    doc.add_paragraph(patient_file.diagnostic or 'Nu este completat')
    
    doc.add_heading('Tratament', level=1)
    doc.add_paragraph(patient_file.tratament or 'Nu este completat')
    
    doc.add_heading('Recomandări', level=1)
    doc.add_paragraph(patient_file.recomandari or 'Nu este completat')
    
    doc.add_heading('Observații', level=1)
    doc.add_paragraph(patient_file.observatii or 'Nu este completat')
    
    # Save document
    filename = f"fisa_pacient_{patient_file.patient_id}_{patient_file.date.strftime('%Y%m%d')}.docx"
    filepath = f"exports/{filename}"
    
    # Create exports directory if it doesn't exist
    os.makedirs("exports", exist_ok=True)
    
    doc.save(filepath)
    
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
